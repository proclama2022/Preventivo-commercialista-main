from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import io
from datetime import datetime

def genera_preventivo(cliente_info, studio_info, tipo_contabilita, servizi):
    doc = Document()
    
    # Aggiungi intestazione
    doc.add_heading('Preventivo', 0)
    
    # Aggiungi informazioni del cliente
    doc.add_heading('Informazioni Cliente', level=1)
    for key, value in cliente_info.items():
        doc.add_paragraph(f"{key.capitalize()}: {value}")
    
    # Aggiungi informazioni dello studio
    doc.add_heading('Informazioni Studio', level=1)
    for key, value in studio_info.items():
        doc.add_paragraph(f"{key.capitalize()}: {value}")
    
    # Aggiungi tipo di contabilità
    doc.add_heading('Tipo di Contabilità', level=1)
    doc.add_paragraph(tipo_contabilita)
    
    # Crea la tabella dei servizi
    doc.add_heading('Dettaglio Servizi', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Descrizione'
    hdr_cells[1].text = 'Quantità'
    hdr_cells[2].text = 'Prezzo'

    # Aggiungi il prezzo base alla tabella dei servizi
    for descrizione, prezzo in servizi['Prezzo Base'].items():
        row_cells = table.add_row().cells
        row_cells[0].text = descrizione
        row_cells[1].text = "1"
        row_cells[2].text = f"€ {prezzo:.2f}"

    # Aggiungi i servizi aggiuntivi
    for servizio, costo in servizi['Servizi Aggiuntivi'].items():
        row_cells = table.add_row().cells
        row_cells[0].text = servizio
        row_cells[1].text = "1"
        row_cells[2].text = f"€ {costo:.2f}"

    # Aggiungi la dichiarazione dei redditi se presente
    if servizi['Dichiarazione dei Redditi'] > 0:
        row_cells = table.add_row().cells
        row_cells[0].text = "Dichiarazione dei Redditi"
        row_cells[1].text = "1"
        row_cells[2].text = f"€ {servizi['Dichiarazione dei Redditi']:.2f}"

    # Calcola il totale
    totale = sum(servizi['Prezzo Base'].values()) + sum(servizi['Servizi Aggiuntivi'].values()) + servizi['Dichiarazione dei Redditi']
    
    # Aggiungi il totale alla tabella
    row_cells = table.add_row().cells
    row_cells[0].text = "Totale"
    row_cells[1].text = ""
    row_cells[2].text = f"€ {totale:.2f}"

    # Salva il documento in memoria
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    
    return doc_stream.getvalue()